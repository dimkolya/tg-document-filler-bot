import aiogram
from aiogram.filters import CommandStart, Command
from aiogram.types import ReplyKeyboardMarkup, KeyboardButton, Message, InlineKeyboardMarkup, InlineKeyboardButton, InputFile
from aiogram.types.input_file import FSInputFile
from docx import Document
import app.user_data
import app.db

import os

from aiogram.fsm.state import StatesGroup, State
from aiogram.fsm.context import FSMContext


class Reg(StatesGroup):
    main = State()
    fill = State()
    filling = State()
    userfilling = State()
    dbupd = State()


directory = 'app/Documents/'

files = list()
for file in os.listdir(directory):
    files.append(file)

router = aiogram.Router()

main = ReplyKeyboardMarkup(keyboard=[
    [KeyboardButton(text="Заполнить документ")]
], resize_keyboard=True)

fill = ReplyKeyboardMarkup(keyboard=[
    [KeyboardButton(text=file)] for file in files
], resize_keyboard=True)


@router.message(aiogram.filters.CommandStart())
@router.message(Reg.main)
async def cmd_start(message: Message, state: FSMContext):
    await state.set_state(Reg.fill)
    await message.answer("Выберите из списка, что хотите сделать", reply_markup=main)


@router.message(Reg.fill)
async def state_fill(message: Message, state: FSMContext):
    await state.set_state(Reg.filling)
    await message.answer("Выберите необходимый документ из списка", reply_markup=fill)


@router.message(Reg.filling)
async def filler(message: Message, state: FSMContext):
    dependencies_rus = list()
    dependencies = list()
    temp = (await state.get_data())
    file = ''
    if "file" not in temp:
        file = message.text
        await state.update_data(file=message.text)
    else:
        file = temp["file"]
        type = app.user_data.simple[(await state.get_data())["main"]]
        info = message.text
        app.db.update(message.from_user.username, {type: info})
        await state.set_state(Reg.filling)
    doc = Document(directory + file)
    for paragraph in doc.paragraphs:
        for rus, eng in app.user_data.simple.items():
            if '[' + eng + ']' in paragraph.text:
                dependencies_rus.append(rus)
                dependencies.append(eng)

    info = app.db.get(message.from_user.username, dependencies)
    s = "Необходимы следующие данные:\n"
    for i in range(len(dependencies)):
        s += dependencies_rus[i] + ":\n\t" + ('' if info[i] is None else info[i]) + '\n'
    s += 'Выберите, какие данные заполнить/изменить. Если все данные верны, нажмите кнопку "ЗАПОЛНИТЬ".'
    dependencies_rus.insert(0, 'ЗАПОЛНИТЬ документ')
    dep_keyboard = ReplyKeyboardMarkup(
        keyboard=[[KeyboardButton(text=rus)] for rus in dependencies_rus],
        resize_keyboard=True
    )
    await message.answer(s, reply_markup=dep_keyboard)
    await state.set_state(Reg.userfilling)


@router.message(Reg.userfilling)
async def filling(message: Message, state: FSMContext):
    if message.text == 'ЗАПОЛНИТЬ документ':
        file = (await state.get_data())["file"]
        doc = Document(directory + file)
        keys = list()
        for rus, eng in app.user_data.simple.items():
            keys.append(eng)
        info = app.db.get(message.from_user.username, keys)
        for paragraph in doc.paragraphs:
            for i in range(len(keys)):
                eng = keys[i]
                if "[" + eng + "]" in paragraph.text:
                    paragraph.text = paragraph.text.replace("[" + eng + "]", info[i])

        # Сохраняем изменения в новый файл

        doc.save("1.docx")
        doc = FSInputFile("1.docx")

        await message.answer_document(doc)

        await state.clear()
        await state.set_state(Reg.fill)
        await message.answer("Выберите из списка, что хотите сделать", reply_markup=main)
        return

    await state.update_data(main=message.text)
    await message.answer("Введите '" + message.text + "'", reply_markup=None)
    await state.set_state(Reg.filling)
